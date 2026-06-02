"""Convert a markdown document to a styled .docx.

Usage: python scripts/md_to_docx.py <source.md> [output.docx]
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


SRC = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("input.md")
DST = Path(sys.argv[2]) if len(sys.argv) > 2 else SRC.with_suffix(".docx")

BODY_FONT = "Calibri"
HEADING_FONT = "Calibri"
BODY_SIZE = Pt(11)
H1_SIZE = Pt(22)
H2_SIZE = Pt(16)
H3_SIZE = Pt(13)


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)

    for name, size, color in (
        ("Heading 1", H1_SIZE, RGBColor(0x1F, 0x3A, 0x5F)),
        ("Heading 2", H2_SIZE, RGBColor(0x1F, 0x3A, 0x5F)),
        ("Heading 3", H3_SIZE, RGBColor(0x2C, 0x5A, 0x8A)),
    ):
        style = doc.styles[name]
        style.font.name = HEADING_FONT
        style.font.size = size
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
INLINE_ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
INLINE_CODE = re.compile(r"`([^`]+)`")


def add_runs(paragraph, text: str) -> None:
    """Tokenize a markdown line for bold/italic/code and emit Word runs."""
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)"
    )
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.text = ""
            paragraph = cell.paragraphs[0]
            add_runs(paragraph, cell_text.strip())
            if r_idx == 0:
                for run in paragraph.runs:
                    run.bold = True
                set_cell_shading(cell, "1F3A5F")
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()


def convert(src: Path, dst: Path) -> None:
    doc = Document()
    configure_styles(doc)

    lines = src.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line:
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        if line.strip() == "---":
            p = doc.add_paragraph()
            p_pr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "999999")
            pbdr.append(bottom)
            p_pr.append(pbdr)
            i += 1
            continue

        if line.lstrip().startswith("|") and "|" in line[1:]:
            table_lines: list[str] = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows: list[list[str]] = []
            for tl in table_lines:
                if re.match(r"^\|[\s\-:|]+\|$", tl):
                    continue
                cells = [c.strip() for c in tl.strip("|").split("|")]
                rows.append(cells)
            add_table(doc, rows)
            continue

        bullet_match = re.match(r"^(\s*)- (.+)$", line)
        number_match = re.match(r"^(\s*)(\d+)\. (.+)$", line)
        if bullet_match:
            indent = len(bullet_match.group(1)) // 2
            style = "List Bullet" if indent == 0 else "List Bullet 2"
            p = doc.add_paragraph(style=style)
            add_runs(p, bullet_match.group(2))
            i += 1
            continue
        if number_match:
            indent = len(number_match.group(1)) // 2
            style = "List Number" if indent == 0 else "List Number 2"
            p = doc.add_paragraph(style=style)
            add_runs(p, number_match.group(3))
            i += 1
            continue

        para_lines = [line]
        i += 1
        while (
            i < len(lines)
            and lines[i].strip()
            and not lines[i].startswith("#")
            and not lines[i].lstrip().startswith("- ")
            and not re.match(r"^\s*\d+\. ", lines[i])
            and not lines[i].lstrip().startswith("|")
            and lines[i].strip() != "---"
        ):
            para_lines.append(lines[i].rstrip())
            i += 1
        text = " ".join(para_lines)
        p = doc.add_paragraph()
        add_runs(p, text)

    doc.save(str(dst))
    print(f"wrote {dst} ({dst.stat().st_size:,} bytes)")


if __name__ == "__main__":
    convert(SRC, DST)
